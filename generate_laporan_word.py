"""
Script untuk menghasilkan Laporan Proyek Akhir dalam format Word (.docx)
yang sesuai dengan standar akademis Indonesia.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ===== PENGATURAN HALAMAN (Standar Akademis Indonesia) =====
for section in doc.sections:
    section.top_margin = Cm(4)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

# ===== PENGATURAN STYLE DEFAULT =====
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.5

# Heading 1 (BAB)
h1_style = doc.styles['Heading 1']
h1_font = h1_style.font
h1_font.name = 'Times New Roman'
h1_font.size = Pt(14)
h1_font.bold = True
h1_font.color.rgb = RGBColor(0, 0, 0)
h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
h1_style.paragraph_format.space_before = Pt(0)
h1_style.paragraph_format.space_after = Pt(12)
h1_style.paragraph_format.line_spacing = 1.5

# Heading 2 (Sub-bab)
h2_style = doc.styles['Heading 2']
h2_font = h2_style.font
h2_font.name = 'Times New Roman'
h2_font.size = Pt(12)
h2_font.bold = True
h2_font.color.rgb = RGBColor(0, 0, 0)
h2_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
h2_style.paragraph_format.space_before = Pt(12)
h2_style.paragraph_format.space_after = Pt(6)
h2_style.paragraph_format.line_spacing = 1.5

# Heading 3 (Sub-sub-bab)
h3_style = doc.styles['Heading 3']
h3_font = h3_style.font
h3_font.name = 'Times New Roman'
h3_font.size = Pt(12)
h3_font.bold = True
h3_font.color.rgb = RGBColor(0, 0, 0)
h3_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
h3_style.paragraph_format.space_before = Pt(12)
h3_style.paragraph_format.space_after = Pt(6)
h3_style.paragraph_format.line_spacing = 1.5

def add_paragraph(text, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, spacing_after=Pt(6), first_line_indent=Cm(1.25)):
    """Menambahkan paragraf standar akademis."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = spacing_after
    p.paragraph_format.line_spacing = 1.5
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return p

def add_mixed_paragraph(parts, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=Cm(1.25), spacing_after=Pt(6)):
    """Menambahkan paragraf dengan campuran bold/italic."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = spacing_after
    p.paragraph_format.line_spacing = 1.5
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = bold
        run.italic = italic
    return p

def add_numbered_item(number, text_parts, indent=Cm(1.25)):
    """Menambahkan item bernomor tanpa first line indent."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = indent
    p.paragraph_format.first_line_indent = Cm(-0.75)
    run = p.add_run(f"{number}. ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    for text, bold, italic in text_parts:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = bold
        run.italic = italic
    return p

def add_bullet_item(text_parts, indent=Cm(1.25)):
    """Menambahkan item bullet."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = indent
    p.paragraph_format.first_line_indent = Cm(-0.5)
    run = p.add_run("• ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    for text, bold, italic in text_parts:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = bold
        run.italic = italic
    return p

def add_table(headers, rows):
    """Menambahkan tabel dengan format standar."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.bold = True
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
    doc.add_paragraph()  # spasi setelah tabel
    return table

def add_empty_lines(count=1):
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

# =====================================================
# HALAMAN JUDUL
# =====================================================
add_empty_lines(3)
add_paragraph("LAPORAN PROYEK AKHIR", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None, spacing_after=Pt(12))
add_empty_lines(1)

title_lines = [
    "IMPLEMENTASI ALGORITMA RANDOM FOREST",
    "DAN K-MEANS CLUSTERING UNTUK PREDIKSI",
    "RISIKO PENYAKIT JANTUNG BERBASIS",
    "APLIKASI WEB"
]
for line in title_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(line)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True

add_empty_lines(3)
add_paragraph("Disusun oleh:", alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
add_empty_lines(1)
add_paragraph("DHARMA MUDITA", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
add_paragraph("NIM: 23312067", alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
add_empty_lines(3)
add_paragraph("Mata Kuliah: Data Mining", alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
add_paragraph("Semester 6 (SP)", alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
add_empty_lines(3)
add_paragraph("PROGRAM STUDI SISTEM INFORMASI", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
add_paragraph("2026", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)

# =====================================================
# HALAMAN DAFTAR ISI
# =====================================================
doc.add_page_break()
doc.add_heading('DAFTAR ISI', level=1)
add_empty_lines(1)

daftar_isi = [
    ("HALAMAN JUDUL", "i"),
    ("DAFTAR ISI", "ii"),
    ("DAFTAR TABEL", "iv"),
    ("DAFTAR GAMBAR", "v"),
    ("", ""),
    ("BAB I    PENDAHULUAN", "1"),
    ("   1.1  Latar Belakang", "1"),
    ("   1.2  Rumusan Masalah", "3"),
    ("   1.3  Tujuan Penelitian", "4"),
    ("   1.4  Manfaat Penelitian", "5"),
    ("   1.5  Batasan Masalah", "6"),
    ("", ""),
    ("BAB II   LANDASAN TEORI", "7"),
    ("   2.1  Data Mining", "7"),
    ("   2.2  Knowledge Discovery in Databases (KDD)", "8"),
    ("   2.3  Algoritma Random Forest", "9"),
    ("   2.4  Algoritma K-Means Clustering", "11"),
    ("   2.5  Evaluasi Model Klasifikasi", "12"),
    ("   2.6  Pra-pemrosesan Data", "13"),
    ("   2.7  Teknologi Pengembangan Aplikasi Web", "14"),
    ("   2.8  Penelitian Terdahulu", "16"),
    ("", ""),
    ("BAB III  METODOLOGI PENELITIAN", "17"),
    ("   3.1  Jenis dan Sumber Data", "17"),
    ("   3.2  Variabel Penelitian", "18"),
    ("   3.3  Alur Penelitian (Tahapan KDD)", "21"),
    ("   3.4  Lingkungan Pengembangan", "22"),
    ("   3.5  Arsitektur Sistem", "23"),
    ("   3.6  Perancangan Antarmuka Pengguna", "24"),
    ("", ""),
    ("BAB IV   HASIL DAN PEMBAHASAN", "25"),
    ("   4.1  Pemilihan Data (Data Selection)", "25"),
    ("   4.2  Analisis Data Eksploratif (EDA)", "26"),
    ("   4.3  Pra-pemrosesan Data (Preprocessing)", "28"),
    ("   4.4  Pemodelan Klasifikasi (Random Forest)", "30"),
    ("   4.5  Evaluasi dan Validasi Model", "31"),
    ("   4.6  Pemodelan Klasterisasi (K-Means)", "33"),
    ("   4.7  Implementasi Aplikasi Web (Deployment)", "34"),
    ("   4.8  Pengujian Fungsional Aplikasi", "37"),
    ("", ""),
    ("BAB V    PENUTUP", "38"),
    ("   5.1  Kesimpulan", "38"),
    ("   5.2  Saran", "39"),
    ("", ""),
    ("DAFTAR PUSTAKA", "41"),
]

for item, page in daftar_isi:
    if item == "":
        add_empty_lines(1)
        continue
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.5
    
    is_bab = item.startswith("BAB") or item in ["HALAMAN JUDUL", "DAFTAR ISI", "DAFTAR TABEL", "DAFTAR GAMBAR", "DAFTAR PUSTAKA"]
    
    run = p.add_run(item)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = is_bab
    
    # Tab dan nomor halaman
    run2 = p.add_run(f"\t{page}")
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)

# =====================================================
# BAB I - PENDAHULUAN
# =====================================================
doc.add_page_break()
doc.add_heading('BAB I', level=1)
doc.add_heading('PENDAHULUAN', level=1)

doc.add_heading('1.1 Latar Belakang', level=2)

add_mixed_paragraph([
    ("Penyakit kardiovaskular atau penyakit jantung (", False, False),
    ("Cardiovascular Disease", False, True),
    (") hingga saat ini masih menduduki peringkat pertama sebagai penyebab kematian tertinggi di dunia. Menurut data resmi ", False, False),
    ("World Health Organization", False, True),
    (" (WHO), diperkirakan sebanyak 17,9 juta jiwa meninggal dunia setiap tahunnya akibat penyakit ini, yang merepresentasikan sekitar 32% dari seluruh kematian global (WHO, 2021). Angka yang sangat mengkhawatirkan ini diperburuk oleh fakta bahwa mayoritas kasus serangan jantung terjadi secara tiba-tiba tanpa gejala awal yang disadari oleh penderitanya, sehingga penanganan medis seringkali terlambat diberikan.", False, False),
])

add_mixed_paragraph([
    ("Di sisi lain, pesatnya perkembangan teknologi informasi dan ilmu komputasi telah melahirkan sebuah cabang keilmuan yang sangat potensial, yaitu ", False, False),
    ("Data Mining", False, True),
    (". ", False, False),
    ("Data Mining", False, True),
    (" merupakan proses mengekstraksi pola-pola tersembunyi yang bernilai dari sekumpulan data berukuran masif (Han, Kamber, & Pei, 2012). Dalam konteks dunia medis, miliaran rekam medis elektronik yang tersimpan di berbagai basis data rumah sakit dan lembaga kesehatan di seluruh dunia kini dapat diolah secara komputasional untuk menemukan korelasi dan pola-pola penyebab penyakit yang tidak terlihat oleh mata manusia.", False, False),
])

add_mixed_paragraph([
    ("Salah satu sumber data medis yang sangat kredibel dan berskala besar adalah ", False, False),
    ("Behavioral Risk Factor Surveillance System", False, True),
    (" (BRFSS) yang dikelola oleh ", False, False),
    ("Centers for Disease Control and Prevention", False, True),
    (" (CDC) Amerika Serikat. Dataset BRFSS tahun 2015 menyimpan lebih dari 253.680 baris data responden yang mencakup informasi komprehensif terkait status kesehatan klinis, perilaku gaya hidup sehari-hari, hingga kondisi sosioekonomi setiap individu (CDC, 2015). Volume dan variasi data sebesar ini menjadikannya kandidat yang sangat ideal untuk dianalisis menggunakan pendekatan ", False, False),
    ("Machine Learning", False, True),
    (".", False, False),
])

add_mixed_paragraph([
    ("Berbeda dengan pendekatan prediksi konvensional yang umumnya hanya mempertimbangkan faktor klinis seperti tekanan darah dan kadar kolesterol, penelitian ini dirancang secara khusus untuk turut mengintegrasikan faktor sosioekonomi, yakni ", False, False),
    ("tingkat pendidikan", True, False),
    (" dan ", False, False),
    ("pendapatan", True, False),
    (", sebagai variabel prediktor. Berdasarkan temuan medis modern, faktor sosioekonomi memiliki dampak langsung yang signifikan terhadap literasi kesehatan seseorang, kualitas gizi harian, tingkat stres akibat tekanan pekerjaan, serta aksesibilitas terhadap fasilitas pelayanan kesehatan dan jaminan asuransi (Braveman, Egerter, & Williams, 2011). Oleh karena itu, penyertaan variabel-variabel ini diharapkan menghasilkan model prediksi yang lebih holistik dan mendekati kondisi dunia nyata.", False, False),
])

add_mixed_paragraph([
    ("Pada puncaknya, sebuah model kecerdasan buatan (", False, False),
    ("Artificial Intelligence", False, True),
    (") sehebat apa pun secara intrinsik akan tetap bersifat teoretis apabila hanya terkurung dalam lingkungan pemrograman mentah seperti Google Colab. Tenaga medis profesional, klinik kesehatan, maupun masyarakat awam membutuhkan perangkat yang dapat diakses dengan mudah dan intuitif. Oleh karena itu, proyek akhir ini tidak hanya berhenti pada tahap pembuatan model, melainkan melangkah lebih jauh dengan mengimplementasikan (", False, False),
    ("deployment", False, True),
    (") algoritma ", False, False),
    ("Random Forest", False, True),
    (" yang telah terlatih ke dalam sebuah ", False, False),
    ("aplikasi web interaktif", True, False),
    (" berbasis ", False, False),
    ("framework", False, True),
    (" Streamlit. Aplikasi web ini dirancang sebagai jembatan penghubung antara kompleksitas kecerdasan buatan dan kebutuhan medis praktis, sehingga siapa saja dapat melakukan asesmen risiko penyakit jantung secara ", False, False),
    ("real-time", False, True),
    (" tanpa memerlukan pengetahuan pemrograman sama sekali.", False, False),
])

# 1.2 Rumusan Masalah
doc.add_heading('1.2 Rumusan Masalah', level=2)
add_paragraph("Berdasarkan latar belakang yang telah diuraikan di atas, maka rumusan masalah dalam proyek akhir ini adalah sebagai berikut:")

add_numbered_item(1, [
    ("Bagaimana proses pra-pemrosesan data (", False, False),
    ("data preprocessing", False, True),
    (") terhadap dataset CDC BRFSS 2015 yang berskala masif (253.680 baris data dan 22 kolom) agar optimal untuk dipelajari oleh algoritma ", False, False),
    ("Machine Learning", False, True),
    ("?", False, False),
])
add_numbered_item(2, [
    ("Bagaimana kinerja dan tingkat akurasi algoritma ", False, False),
    ("Random Forest Classifier", False, True),
    (" dalam memprediksi risiko penyakit jantung berdasarkan 21 fitur prediktor yang mencakup aspek klinis, gaya hidup, dan sosioekonomi?", False, False),
])
add_numbered_item(3, [
    ("Bagaimana mengidentifikasi segmentasi atau pengelompokan profil pasien secara alami menggunakan pendekatan ", False, False),
    ("Unsupervised Learning", False, True),
    (" melalui algoritma klasterisasi ", False, False),
    ("K-Means", False, True),
    (" pada dimensi Umur dan Indeks Massa Tubuh (BMI)?", False, False),
])
add_numbered_item(4, [
    ("Bagaimana cara mengintegrasikan model prediksi yang telah dilatih dan disimpan dalam format ", False, False),
    ("pickle", False, True),
    (" (.pkl) ke dalam sistem aplikasi web fungsional berbasis ", False, False),
    ("framework", False, True),
    (" Streamlit yang siap digunakan secara ", False, False),
    ("real-time", False, True),
    ("?", False, False),
])

# 1.3 Tujuan Penelitian
doc.add_heading('1.3 Tujuan Penelitian', level=2)
add_paragraph("Adapun tujuan dari pelaksanaan proyek akhir ini adalah sebagai berikut:")

add_numbered_item(1, [
    ("Melakukan rekayasa dan standardisasi data pada skala besar, termasuk penanganan dan mitigasi terhadap kondisi data yang tidak seimbang (", False, False),
    ("imbalanced data", False, True),
    ("), sehingga dataset siap diolah oleh mesin secara optimal.", False, False),
])
add_numbered_item(2, [
    ("Membangun model ", False, False),
    ("Machine Learning", False, True),
    (" berbasis algoritma ", False, False),
    ("Random Forest", False, True),
    (" yang mampu mendeteksi potensi risiko penyakit jantung dengan tingkat akurasi evaluasi yang teruji secara statistik (target di atas 85%).", False, False),
])
add_numbered_item(3, [
    ("Menganalisis penyebaran klaster atau segmen pasien menggunakan algoritma ", False, False),
    ("K-Means", False, True),
    (" untuk membuktikan bahwa pola-pola gaya hidup pasien dapat dipetakan dan dikelompokkan secara akurat tanpa campur tangan label manual manusia.", False, False),
])
add_numbered_item(4, [
    ("Menciptakan produk akhir perangkat lunak (", False, False),
    ("software", False, True),
    (") berbentuk aplikasi web interaktif yang ramah pengguna (", False, False),
    ("user-friendly", False, True),
    ("), mampu menerima masukan data pasien, dan menghasilkan prediksi risiko beserta rekomendasi medis secara otomatis dan ", False, False),
    ("real-time", False, True),
    (".", False, False),
])

# 1.4 Manfaat Penelitian
doc.add_heading('1.4 Manfaat Penelitian', level=2)
add_paragraph("Manfaat yang diharapkan dari pelaksanaan proyek akhir ini terbagi menjadi dua kategori utama, yaitu:")

doc.add_heading('1.4.1 Manfaat Teoretis', level=3)
add_numbered_item(1, [("Memberikan kontribusi ilmiah terhadap pemahaman penerapan algoritma ensemble learning (Random Forest) dan unsupervised learning (K-Means) pada domain kesehatan masyarakat, khususnya dalam prediksi risiko penyakit jantung.", False, False)])
add_numbered_item(2, [("Memperkaya referensi akademis tentang bagaimana faktor sosioekonomi (pendidikan dan pendapatan) dapat diintegrasikan sebagai variabel prediktor yang bermakna dalam model prediksi penyakit, di samping faktor klinis konvensional.", False, False)])
add_numbered_item(3, [("Memberikan bukti empiris bahwa dataset berskala besar (lebih dari 250.000 rekam medis) dari lembaga kredibel internasional (CDC) dapat diolah secara efektif menggunakan pendekatan Data Mining untuk menghasilkan wawasan medis yang bernilai tinggi.", False, False)])

doc.add_heading('1.4.2 Manfaat Praktis', level=3)
add_numbered_item(1, [
    ("Bagi Tenaga Medis: ", True, False),
    ("Menyediakan alat bantu skrining awal (screening tool) yang cepat dan efisien untuk mengidentifikasi pasien dengan risiko tinggi penyakit jantung, sehingga prioritas penanganan medis dapat dilakukan lebih dini.", False, False),
])
add_numbered_item(2, [
    ("Bagi Masyarakat Umum: ", True, False),
    ("Menyediakan platform yang mudah diakses dan dioperasikan untuk melakukan asesmen mandiri terhadap risiko kesehatan jantung berdasarkan kondisi personal masing-masing individu.", False, False),
])
add_numbered_item(3, [
    ("Bagi Institusi Pendidikan: ", True, False),
    ("Menjadi referensi dan studi kasus nyata tentang implementasi end-to-end dari proyek Data Mining, mulai dari eksplorasi data mentah, pemodelan algoritma, hingga deployment ke dalam produk perangkat lunak yang fungsional.", False, False),
])
add_numbered_item(4, [
    ("Bagi Pengembang Sistem: ", True, False),
    ("Mendemonstrasikan arsitektur integrasi antara model Machine Learning (file .pkl) dengan framework aplikasi web (Streamlit) sebagai pola desain (design pattern) yang dapat direplikasi untuk proyek serupa di domain lain.", False, False),
])

# 1.5 Batasan Masalah
doc.add_heading('1.5 Batasan Masalah', level=2)
add_paragraph("Agar pembahasan dalam proyek akhir ini tetap terfokus dan terarah, maka ditetapkan batasan-batasan masalah sebagai berikut:")

add_numbered_item(1, [("Dataset yang digunakan adalah data sekunder dari CDC BRFSS tahun 2015 yang bersifat statis dan tidak diperbarui secara berkala (real-time).", False, False)])
add_numbered_item(2, [("Algoritma klasifikasi yang digunakan terbatas pada Random Forest Classifier, dan algoritma klasterisasi terbatas pada K-Means Clustering. Perbandingan dengan algoritma lain (seperti Support Vector Machine, Gradient Boosting, atau DBSCAN) tidak termasuk dalam lingkup penelitian ini.", False, False)])
add_numbered_item(3, [("Aplikasi web yang dikembangkan menggunakan framework Streamlit berjalan secara lokal (localhost) dan belum di-deploy ke server produksi berbasis cloud (seperti Heroku atau AWS).", False, False)])
add_numbered_item(4, [("Hasil prediksi yang dihasilkan oleh sistem bersifat informatif sebagai alat bantu skrining awal dan ", False, False), ("bukan", True, False), (" merupakan pengganti diagnosis medis resmi dari dokter profesional.", False, False)])
add_numbered_item(5, [("Fitur-fitur yang digunakan sebagai variabel prediktor dibatasi pada 21 kolom yang tersedia dalam dataset CDC BRFSS 2015, tanpa penambahan fitur eksternal dari sumber data lain.", False, False)])

# =====================================================
# BAB II - LANDASAN TEORI
# =====================================================
doc.add_page_break()
doc.add_heading('BAB II', level=1)
doc.add_heading('LANDASAN TEORI', level=1)

# 2.1 Data Mining
doc.add_heading('2.1 Data Mining', level=2)
add_mixed_paragraph([
    ("Data Mining", False, True),
    (" adalah proses penemuan pola-pola yang menarik dan berpotensi berguna dari sekumpulan data berukuran besar (", False, False),
    ("large datasets", False, True),
    ("). Menurut Han, Kamber, dan Pei (2012) dalam buku ", False, False),
    ("Data Mining: Concepts and Techniques", False, True),
    (", ", False, False),
    ("Data Mining", False, True),
    (" didefinisikan sebagai proses mengekstraksi atau \"menambang\" pengetahuan (", False, False),
    ("knowledge", False, True),
    (") dari volume data yang masif. Proses ini melibatkan serangkaian teknik yang berasal dari irisan beberapa bidang keilmuan, di antaranya statistika, kecerdasan buatan (", False, False),
    ("Artificial Intelligence", False, True),
    ("), ", False, False),
    ("Machine Learning", False, True),
    (", dan pengelolaan basis data.", False, False),
])

add_mixed_paragraph([
    ("Secara fundamental, ", False, False),
    ("Data Mining", False, True),
    (" bertujuan untuk mengubah data mentah yang tampak tidak bermakna menjadi informasi dan pengetahuan yang dapat ditindaklanjuti (", False, False),
    ("actionable knowledge", False, True),
    ("). Dalam konteks penelitian ini, ", False, False),
    ("Data Mining", False, True),
    (" digunakan untuk menemukan pola-pola tersembunyi dalam ratusan ribu rekam medis pasien yang berkaitan dengan risiko penyakit jantung.", False, False),
])

add_paragraph("Tugas-tugas utama dalam Data Mining meliputi (Larose & Larose, 2014):")
add_numbered_item(1, [("Klasifikasi (", True, False), ("Classification", True, True), ("): ", True, False), ("Memprediksi kategori atau label dari sebuah data baru berdasarkan pola yang telah dipelajari dari data historis.", False, False)])
add_numbered_item(2, [("Klasterisasi (", True, False), ("Clustering", True, True), ("): ", True, False), ("Mengelompokkan data ke dalam beberapa segmen berdasarkan kesamaan karakteristik tanpa menggunakan label yang telah ditentukan sebelumnya.", False, False)])
add_numbered_item(3, [("Asosiasi (", True, False), ("Association Rule Mining", True, True), ("): ", True, False), ("Menemukan hubungan antaritem dalam sekumpulan data transaksi.", False, False)])
add_numbered_item(4, [("Regresi (", True, False), ("Regression", True, True), ("): ", True, False), ("Memprediksi nilai numerik kontinu berdasarkan variabel-variabel prediktor.", False, False)])

# 2.2 KDD
doc.add_heading('2.2 Knowledge Discovery in Databases (KDD)', level=2)
add_mixed_paragraph([
    ("Knowledge Discovery in Databases", False, True),
    (" (KDD) merupakan kerangka kerja (", False, False),
    ("framework", False, True),
    (") menyeluruh yang menempatkan ", False, False),
    ("Data Mining", False, True),
    (" sebagai salah satu tahapan intinya. Menurut Fayyad, Piatetsky-Shapiro, dan Smyth (1996), KDD adalah proses nontrivial untuk mengidentifikasi pola-pola yang valid, baru, berpotensi bermanfaat, dan pada akhirnya dapat dipahami di dalam data. Tahapan-tahapan KDD yang menjadi acuan dalam proyek ini meliputi:", False, False),
])

add_numbered_item(1, [("Pemilihan Data (", True, False), ("Data Selection", True, True), ("): ", True, False), ("Mengidentifikasi dan memilih subset data yang relevan dari sumber data yang tersedia.", False, False)])
add_numbered_item(2, [("Pra-pemrosesan Data (", True, False), ("Data Preprocessing", True, True), ("): ", True, False), ("Membersihkan data dari noise, menangani data yang hilang (missing values), serta melakukan transformasi data.", False, False)])
add_numbered_item(3, [("Transformasi Data (", True, False), ("Data Transformation", True, True), ("): ", True, False), ("Mengubah format atau skala data ke dalam bentuk yang optimal untuk diproses oleh algoritma Machine Learning.", False, False)])
add_numbered_item(4, [("Data Mining: ", True, False), ("Menerapkan algoritma-algoritma tertentu (klasifikasi, klasterisasi, dsb.) untuk mengekstraksi pola dari data yang telah ditransformasi.", False, False)])
add_numbered_item(5, [("Evaluasi dan Interpretasi: ", True, False), ("Menilai validitas dan kebermaknaan pola-pola yang ditemukan, serta menerjemahkannya menjadi pengetahuan yang dapat dipahami.", False, False)])

# 2.3 Random Forest
doc.add_heading('2.3 Algoritma Random Forest', level=2)
add_mixed_paragraph([
    ("Random Forest", False, True),
    (" adalah algoritma ", False, False),
    ("Machine Learning", False, True),
    (" berbasis ", False, False),
    ("ensemble learning", False, True),
    (" yang dikembangkan oleh Leo Breiman pada tahun 2001. Algoritma ini bekerja dengan prinsip membangun banyak pohon keputusan (", False, False),
    ("Decision Trees", False, True),
    (") secara simultan pada saat proses pelatihan (", False, False),
    ("training", False, True),
    ("), kemudian menghasilkan prediksi akhir berdasarkan modus (suara terbanyak) dari seluruh pohon tersebut untuk kasus klasifikasi, atau rata-rata untuk kasus regresi (Breiman, 2001).", False, False),
])

doc.add_heading('2.3.1 Mekanisme Kerja Random Forest', level=3)
add_paragraph("Mekanisme kerja Random Forest dapat dijabarkan sebagai berikut:")
add_numbered_item(1, [("Bagging (", True, False), ("Bootstrap Aggregating", True, True), ("): ", True, False), ("Dari dataset pelatihan asli, dibuatkan beberapa subset data secara acak dengan pengembalian (sampling with replacement). Setiap subset data ini akan digunakan untuk melatih satu pohon keputusan yang independen.", False, False)])
add_numbered_item(2, [("Pemilihan Fitur Acak (", True, False), ("Random Feature Selection", True, True), ("): ", True, False), ("Pada setiap titik pemisahan (split) dalam pohon keputusan, hanya sejumlah kecil fitur yang dipilih secara acak untuk dipertimbangkan. Mekanisme ini mengurangi korelasi antar pohon dan meningkatkan keragaman (diversity) model.", False, False)])
add_numbered_item(3, [("Voting (Agregasi): ", True, False), ("Setelah seluruh pohon keputusan selesai dibangun, prediksi akhir ditentukan melalui mekanisme majority voting (suara terbanyak). Pendekatan ini secara signifikan mengurangi risiko overfitting.", False, False)])

doc.add_heading('2.3.2 Keunggulan Random Forest', level=3)
add_paragraph("Menurut Hastie, Tibshirani, dan Friedman (2009), beberapa keunggulan utama Random Forest meliputi:")
add_bullet_item([("Mampu menangani dataset berdimensi tinggi (", False, False), ("high-dimensional data", False, True), (") dengan jumlah fitur yang besar.", False, False)])
add_bullet_item([("Memiliki ketahanan yang sangat baik terhadap ", False, False), ("overfitting", False, True), (" berkat mekanisme ", False, False), ("bagging", False, True), (" dan pemilihan fitur acak.", False, False)])
add_bullet_item([("Dapat memberikan estimasi ", False, False), ("feature importance", False, True), (" (tingkat kepentingan masing-masing fitur) yang berguna untuk interpretasi model.", False, False)])
add_bullet_item([("Mampu menangani data kategorik dan numerik secara bersamaan tanpa memerlukan banyak pra-pemrosesan tambahan.", False, False)])

# 2.4 K-Means
doc.add_heading('2.4 Algoritma K-Means Clustering', level=2)
add_mixed_paragraph([
    ("K-Means Clustering", False, True),
    (" adalah algoritma pengelompokan berbasis partisi (", False, False),
    ("partitional clustering", False, True),
    (") yang bertujuan untuk membagi ", False, False),
    ("n", False, True),
    (" observasi ke dalam ", False, False),
    ("k", False, True),
    (" klaster, di mana setiap observasi ditempatkan ke dalam klaster dengan titik pusat (", False, False),
    ("centroid", False, True),
    (") terdekat (MacQueen, 1967). Algoritma ini termasuk dalam kategori ", False, False),
    ("Unsupervised Learning", False, True),
    (" karena tidak memerlukan label target yang telah ditentukan sebelumnya.", False, False),
])

doc.add_heading('2.4.1 Mekanisme Kerja K-Means', level=3)
add_paragraph("Langkah-langkah kerja algoritma K-Means adalah sebagai berikut:")
add_numbered_item(1, [("Inisialisasi: ", True, False), ("Tentukan jumlah klaster (k) yang diinginkan, kemudian inisialisasi k buah titik pusat (centroid) secara acak di dalam ruang fitur.", False, False)])
add_numbered_item(2, [("Penugasan (", True, False), ("Assignment", True, True), ("): ", True, False), ("Hitung jarak setiap titik data terhadap semua centroid menggunakan metrik jarak Euclidean, lalu tugaskan setiap titik data ke klaster dengan centroid terdekat.", False, False)])
add_numbered_item(3, [("Pembaruan (", True, False), ("Update", True, True), ("): ", True, False), ("Hitung ulang posisi centroid baru untuk setiap klaster berdasarkan rata-rata koordinat seluruh titik data yang menjadi anggotanya.", False, False)])
add_numbered_item(4, [("Iterasi: ", True, False), ("Ulangi langkah 2 dan 3 hingga posisi centroid tidak lagi berubah secara signifikan (konvergen) atau telah mencapai jumlah iterasi maksimum yang ditentukan.", False, False)])

# 2.5 Evaluasi Model
doc.add_heading('2.5 Evaluasi Model Klasifikasi', level=2)
add_paragraph("Untuk menilai kinerja model klasifikasi yang telah dibangun, beberapa metrik evaluasi standar yang digunakan dalam penelitian ini meliputi (Sokolova & Lapalme, 2009):")
add_numbered_item(1, [("Akurasi (", True, False), ("Accuracy", True, True), ("): ", True, False), ("Proporsi prediksi yang benar dari keseluruhan jumlah prediksi.", False, False)])
add_numbered_item(2, [("Presisi (", True, False), ("Precision", True, True), ("): ", True, False), ("Proporsi prediksi positif yang benar-benar positif.", False, False)])
add_numbered_item(3, [("Recall ", True, False), ("(Sensitivitas): ", True, False), ("Proporsi data positif aktual yang berhasil diidentifikasi dengan benar oleh model.", False, False)])
add_numbered_item(4, [("F1-Score: ", True, False), ("Rata-rata harmonik antara Presisi dan Recall, yang memberikan keseimbangan antara keduanya.", False, False)])
add_numbered_item(5, [("Confusion Matrix: ", True, False), ("Tabel matriks yang memvisualisasikan kinerja model dengan merinci empat komponen: True Positive (TP), True Negative (TN), False Positive (FP), dan False Negative (FN).", False, False)])

# 2.6 Pra-pemrosesan
doc.add_heading('2.6 Pra-pemrosesan Data', level=2)
add_paragraph("Pra-pemrosesan data merupakan tahapan krusial yang menentukan kualitas hasil akhir model Machine Learning. Beberapa teknik pra-pemrosesan yang diterapkan dalam penelitian ini meliputi:")

doc.add_heading('2.6.1 Standardisasi (Standard Scaling)', level=3)
add_mixed_paragraph([
    ("Standardisasi atau ", False, False),
    ("Z-Score Normalization", False, True),
    (" adalah teknik transformasi data yang mengubah distribusi setiap fitur sehingga memiliki rata-rata (", False, False),
    ("mean", False, True),
    (") = 0 dan simpangan baku (", False, False),
    ("standard deviation", False, True),
    (") = 1. Rumusnya adalah: z = (x - μ) / σ, di mana x adalah nilai asli, μ adalah rata-rata fitur, dan σ adalah simpangan baku fitur (Zheng & Casari, 2018). Teknik ini sangat penting karena algoritma berbasis jarak (seperti K-Means) akan menghasilkan output yang bias jika skala antarfitur tidak seragam.", False, False),
])

doc.add_heading('2.6.2 Pembagian Data (Data Splitting)', level=3)
add_mixed_paragraph([
    ("Pembagian data ke dalam ", False, False),
    ("training set", False, True),
    (" dan ", False, False),
    ("testing set", False, True),
    (" merupakan praktik standar untuk menghindari ", False, False),
    ("overfitting", False, True),
    (" dan menguji kemampuan generalisasi model. Dalam proyek ini, digunakan rasio pembagian 80:20, yakni 80% data untuk pelatihan dan 20% data untuk pengujian. Parameter stratify=y diterapkan untuk memastikan proporsi distribusi kelas target tetap konsisten di kedua subset (Géron, 2019).", False, False),
])

doc.add_heading('2.6.3 Penanganan Data Tidak Seimbang', level=3)
add_mixed_paragraph([
    ("Data medis pada umumnya bersifat tidak seimbang (", False, False),
    ("imbalanced", False, True),
    ("), di mana jumlah kasus positif (pasien sakit) jauh lebih sedikit dibandingkan kasus negatif (pasien sehat). Untuk mengatasi hal ini, parameter class_weight='balanced' diaktifkan pada model ", False, False),
    ("Random Forest", False, True),
    (", yang secara otomatis memberikan bobot lebih tinggi pada kelas minoritas saat proses pelatihan (Chawla, 2010).", False, False),
])

# 2.7 Teknologi
doc.add_heading('2.7 Teknologi Pengembangan Aplikasi Web', level=2)

doc.add_heading('2.7.1 Python', level=3)
add_mixed_paragraph([
    ("Python adalah bahasa pemrograman tingkat tinggi yang bersifat ", False, False),
    ("open-source", False, True),
    (", interpretatif, dan mendukung berbagai paradigma pemrograman. Python menjadi bahasa yang paling populer dalam bidang ", False, False),
    ("Data Science", False, True),
    (" dan ", False, False),
    ("Machine Learning", False, True),
    (" berkat ekosistem pustaka (", False, False),
    ("library", False, True),
    (") yang sangat kaya, di antaranya Pandas, NumPy, Scikit-Learn, Matplotlib, dan Seaborn (VanderPlas, 2016).", False, False),
])

doc.add_heading('2.7.2 Streamlit', level=3)
add_mixed_paragraph([
    ("Streamlit adalah ", False, False),
    ("framework", False, True),
    (" aplikasi web berbasis Python yang dirancang khusus untuk membuat aplikasi data interaktif dengan cepat dan mudah. Keunggulan utama Streamlit terletak pada kemampuannya untuk mengonversi skrip Python biasa menjadi aplikasi web tanpa memerlukan pengetahuan pengembangan web tradisional seperti HTML, CSS, atau JavaScript secara mendalam (Streamlit Inc., 2024).", False, False),
])

doc.add_heading('2.7.3 Scikit-Learn', level=3)
add_mixed_paragraph([
    ("Scikit-Learn adalah pustaka ", False, False),
    ("Machine Learning", False, True),
    (" untuk Python yang menyediakan implementasi berbagai algoritma klasifikasi, regresi, klasterisasi, dan pra-pemrosesan data. Pustaka ini dibangun di atas NumPy, SciPy, dan Matplotlib, serta mengikuti antarmuka pemrograman yang konsisten dan terdokumentasi dengan sangat baik (Pedregosa et al., 2011).", False, False),
])

doc.add_heading('2.7.4 Joblib', level=3)
add_mixed_paragraph([
    ("Joblib adalah pustaka Python yang digunakan untuk melakukan serialisasi objek-objek Python berukuran besar (seperti model ", False, False),
    ("Machine Learning", False, True),
    (") ke dalam format file .pkl (", False, False),
    ("pickle", False, True),
    ("). File .pkl ini memungkinkan model yang telah dilatih untuk disimpan secara permanen dan dimuat kembali (", False, False),
    ("load", False, True),
    (") di lingkungan yang berbeda, misalnya di dalam aplikasi web, tanpa perlu melatih ulang model dari awal.", False, False),
])

# 2.8 Penelitian Terdahulu
doc.add_heading('2.8 Penelitian Terdahulu', level=2)
add_paragraph("Beberapa penelitian terdahulu yang relevan dan menjadi acuan dalam proyek akhir ini disajikan dalam tabel berikut:")

add_table(
    ["No.", "Peneliti (Tahun)", "Judul", "Algoritma", "Akurasi"],
    [
        ["1", "Dinh et al. (2019)", "A Data-Driven Approach to Predicting Diabetes and Cardiovascular Disease with ML", "Random Forest, Gradient Boosting", "90,2%"],
        ["2", "Alaa et al. (2019)", "Cardiovascular Disease Risk Prediction Using Automated ML", "AutoPrognosis (Ensemble)", "89,1%"],
        ["3", "Mohan et al. (2019)", "Effective Heart Disease Prediction Using Hybrid ML Techniques", "Random Forest + Feature Selection", "88,7%"],
        ["4", "Pal & Parija (2021)", "Prediction of Heart Diseases Using Random Forest", "Random Forest", "87,5%"],
        ["5", "Reddy et al. (2020)", "Heart Disease Risk Prediction Using ML Algorithm", "Random Forest, SVM, KNN", "90,16%"],
    ]
)

add_mixed_paragraph([
    ("Berdasarkan tinjauan penelitian terdahulu di atas, dapat disimpulkan bahwa algoritma ", False, False),
    ("Random Forest", False, True),
    (" secara konsisten menunjukkan kinerja yang unggul dan stabil dalam domain prediksi penyakit jantung dibandingkan algoritma-algoritma lainnya. Hal ini memperkuat keputusan pemilihan ", False, False),
    ("Random Forest", False, True),
    (" sebagai algoritma utama dalam proyek akhir ini.", False, False),
])

# =====================================================
# BAB III - METODOLOGI
# =====================================================
doc.add_page_break()
doc.add_heading('BAB III', level=1)
doc.add_heading('METODOLOGI PENELITIAN', level=1)

doc.add_heading('3.1 Jenis dan Sumber Data', level=2)
add_mixed_paragraph([
    ("Penelitian ini menggunakan ", False, False),
    ("data sekunder", True, False),
    (" yang bersumber dari ", False, False),
    ("Behavioral Risk Factor Surveillance System", False, True),
    (" (BRFSS) tahun 2015, sebuah survei kesehatan tahunan berskala nasional yang diselenggarakan oleh ", False, False),
    ("Centers for Disease Control and Prevention", False, True),
    (" (CDC) Amerika Serikat. Dataset ini tersedia secara publik melalui repositori Kaggle dan telah digunakan secara luas dalam berbagai penelitian akademis di bidang kesehatan masyarakat dan kecerdasan buatan.", False, False),
])

add_paragraph("Spesifikasi dataset yang digunakan adalah sebagai berikut:")
add_bullet_item([("Nama File: ", True, False), ("heart_disease_health_indicators_BRFSS2015.csv", False, False)])
add_bullet_item([("Jumlah Baris Data: ", True, False), ("253.680 baris (rekam medis pasien)", False, False)])
add_bullet_item([("Jumlah Kolom: ", True, False), ("22 kolom (1 kolom target + 21 kolom fitur prediktor)", False, False)])
add_bullet_item([("Ukuran File: ", True, False), ("±22 MB", False, False)])
add_bullet_item([("Jenis Klasifikasi: ", True, False), ("Klasifikasi Biner (Binary Classification)", False, False)])

# 3.2 Variabel Penelitian
doc.add_heading('3.2 Variabel Penelitian', level=2)
add_paragraph("Variabel-variabel yang digunakan dalam penelitian ini dikelompokkan menjadi dua kategori utama:")

doc.add_heading('3.2.1 Variabel Dependen (Target/Label)', level=3)
add_table(
    ["Nama Variabel", "Keterangan", "Nilai"],
    [["HeartDiseaseorAttack", "Riwayat penyakit jantung atau serangan jantung", "0 = Tidak, 1 = Ya"]]
)

doc.add_heading('3.2.2 Variabel Independen (Fitur Prediktor)', level=3)
add_paragraph("Ke-21 variabel independen dikelompokkan ke dalam tiga aspek utama:")

add_paragraph("A. Faktor Klinis (Kondisi Medis)", bold=True, first_line_indent=None)
add_table(
    ["No.", "Nama Variabel", "Keterangan"],
    [
        ["1", "HighBP", "Riwayat tekanan darah tinggi (0/1)"],
        ["2", "HighChol", "Riwayat kolesterol tinggi (0/1)"],
        ["3", "CholCheck", "Pemeriksaan kolesterol dalam 5 tahun terakhir (0/1)"],
        ["4", "BMI", "Indeks Massa Tubuh (numerik kontinu)"],
        ["5", "Stroke", "Riwayat stroke (0/1)"],
        ["6", "Diabetes", "Status diabetes (0/1/2)"],
        ["7", "GenHlth", "Evaluasi kesehatan umum secara pribadi (1-5)"],
        ["8", "MentHlth", "Jumlah hari kesehatan mental memburuk (0-30)"],
        ["9", "PhysHlth", "Jumlah hari kesehatan fisik memburuk (0-30)"],
        ["10", "DiffWalk", "Kesulitan berjalan atau naik tangga (0/1)"],
    ]
)

add_paragraph("B. Faktor Gaya Hidup (Perilaku)", bold=True, first_line_indent=None)
add_table(
    ["No.", "Nama Variabel", "Keterangan"],
    [
        ["11", "Smoker", "Status perokok aktif (0/1)"],
        ["12", "PhysActivity", "Aktivitas fisik dalam 30 hari terakhir (0/1)"],
        ["13", "Fruits", "Konsumsi buah minimal 1x sehari (0/1)"],
        ["14", "Veggies", "Konsumsi sayur minimal 1x sehari (0/1)"],
        ["15", "HvyAlcoholConsump", "Konsumsi alkohol berat (0/1)"],
        ["16", "AnyHealthcare", "Kepemilikan asuransi kesehatan (0/1)"],
        ["17", "NoDocbcCost", "Pernah batal ke dokter karena biaya (0/1)"],
    ]
)

add_paragraph("C. Faktor Sosioekonomi dan Demografi", bold=True, first_line_indent=None)
add_table(
    ["No.", "Nama Variabel", "Keterangan"],
    [
        ["18", "Sex", "Jenis kelamin (0 = Perempuan, 1 = Laki-laki)"],
        ["19", "Age", "Kategori rentang usia (1-13)"],
        ["20", "Education", "Tingkat pendidikan terakhir (1-6)"],
        ["21", "Income", "Tingkat pendapatan (1-8)"],
    ]
)

# 3.3 Alur Penelitian
doc.add_heading('3.3 Alur Penelitian (Tahapan KDD)', level=2)
add_paragraph("Alur penelitian yang diterapkan dalam proyek akhir ini mengikuti kerangka kerja Knowledge Discovery in Databases (KDD) yang terdiri dari enam tahapan utama:")
add_numbered_item(1, [("Tahap 1 – Pemilihan Data (", True, False), ("Data Selection", True, True), ("): ", True, False), ("Memuat dataset CSV ke dalam memori sistem dan melakukan preview awal terhadap struktur data.", False, False)])
add_numbered_item(2, [("Tahap 2 – Analisis Data Eksploratif (EDA): ", True, False), ("Memverifikasi keutuhan data (missing values), menganalisis distribusi kelas target, dan memvisualisasikan korelasi antarfitur melalui heatmap.", False, False)])
add_numbered_item(3, [("Tahap 3 – Pra-pemrosesan Data (", True, False), ("Preprocessing", True, True), ("): ", True, False), ("Memisahkan variabel independen dan dependen, membagi data menjadi training set (80%) dan testing set (20%), serta melakukan standardisasi menggunakan StandardScaler.", False, False)])
add_numbered_item(4, [("Tahap 4 – Pemodelan Klasifikasi: ", True, False), ("Membangun model Random Forest Classifier dengan 100 pohon keputusan (n_estimators=100) dan parameter class_weight='balanced'.", False, False)])
add_numbered_item(5, [("Tahap 5 – Evaluasi dan Validasi: ", True, False), ("Menguji model menggunakan data uji, menghitung akurasi, mencetak Classification Report, dan memvisualisasikan Confusion Matrix.", False, False)])
add_numbered_item(6, [("Tahap 6 – Klasterisasi dan Deployment: ", True, False), ("Menerapkan K-Means Clustering untuk segmentasi pasien, mengekspor model ke format .pkl, dan mengintegrasikannya ke dalam aplikasi web Streamlit.", False, False)])

# 3.4 Lingkungan Pengembangan
doc.add_heading('3.4 Lingkungan Pengembangan', level=2)
add_paragraph("Perangkat keras dan perangkat lunak yang digunakan dalam pengembangan proyek akhir ini adalah sebagai berikut:")
add_table(
    ["Komponen", "Spesifikasi"],
    [
        ["Sistem Operasi", "Windows"],
        ["Bahasa Pemrograman", "Python 3.10"],
        ["Lingkungan Notebook", "Google Colaboratory (Colab)"],
        ["Lingkungan Web", "Visual Studio Code"],
        ["Framework Aplikasi Web", "Streamlit"],
        ["Pustaka Machine Learning", "Scikit-Learn"],
        ["Pustaka Manipulasi Data", "Pandas, NumPy"],
        ["Pustaka Visualisasi", "Matplotlib, Seaborn"],
        ["Pustaka Serialisasi Model", "Joblib"],
    ]
)

# 3.5 Arsitektur Sistem
doc.add_heading('3.5 Arsitektur Sistem', level=2)
add_paragraph("Arsitektur sistem dalam proyek akhir ini terbagi menjadi dua komponen utama yang saling terintegrasi:")
add_numbered_item(1, [("Komponen Backend (Google Colab): ", True, False), ("Bertanggung jawab atas seluruh proses Data Mining, mulai dari pemilihan data, EDA, pra-pemrosesan, pemodelan klasifikasi (Random Forest), evaluasi, hingga klasterisasi (K-Means). Hasil akhir dari komponen ini adalah dua buah file model terlatih, yaitu model_jantung_rf.pkl (model Random Forest) dan scaler_jantung.pkl (objek StandardScaler).", False, False)])
add_numbered_item(2, [("Komponen Frontend (Streamlit): ", True, False), ("Bertanggung jawab atas penyajian antarmuka pengguna yang interaktif. Komponen ini memuat (load) kedua file .pkl tersebut, menerima masukan data dari pengguna melalui formulir bertahap (wizard-style), melakukan transformasi data menggunakan scaler yang sama, lalu meneruskan data tersebut ke model Random Forest untuk menghasilkan prediksi secara real-time.", False, False)])

# 3.6 Perancangan UI
doc.add_heading('3.6 Perancangan Antarmuka Pengguna', level=2)
add_paragraph("Antarmuka pengguna aplikasi web dirancang dengan pendekatan wizard-style (bertahap) yang membagi 21 variabel masukan ke dalam empat langkah sistematis:")
add_numbered_item(1, [("Langkah 1 – Informasi Dasar Pasien: ", True, False), ("Menampung data demografi dasar meliputi rentang usia, jenis kelamin, tingkat pendidikan terakhir, dan estimasi pendapatan per bulan.", False, False)])
add_numbered_item(2, [("Langkah 2 – Kondisi dan Riwayat Penyakit: ", True, False), ("Menampung data klinis meliputi riwayat tekanan darah tinggi, kolesterol tinggi, pemeriksaan kolesterol, Indeks Massa Tubuh (BMI), riwayat stroke, dan status diabetes.", False, False)])
add_numbered_item(3, [("Langkah 3 – Kebiasaan dan Gaya Hidup: ", True, False), ("Menampung data perilaku meliputi status perokok, aktivitas fisik, konsumsi alkohol, konsumsi buah, dan konsumsi sayur.", False, False)])
add_numbered_item(4, [("Langkah 4 – Keluhan Saat Ini dan Prediksi: ", True, False), ("Menampung data keluhan fisik dan mental terkini, lalu menjalankan prediksi melalui model Random Forest. Hasil prediksi disajikan dalam bentuk klasifikasi (Sehat/Berisiko), persentase keyakinan algoritma (probability), dan rekomendasi medis otomatis.", False, False)])

add_paragraph("Desain bertahap ini diterapkan secara sengaja untuk mengurangi beban kognitif (cognitive load) pengguna, sehingga proses pengisian data terasa lebih ringan, sistematis, dan tidak memusingkan dibandingkan dengan formulir satu halaman penuh yang menampilkan seluruh 21 pertanyaan secara bersamaan.")

# =====================================================
# BAB IV - HASIL DAN PEMBAHASAN
# =====================================================
doc.add_page_break()
doc.add_heading('BAB IV', level=1)
doc.add_heading('HASIL DAN PEMBAHASAN', level=1)

doc.add_heading('4.1 Pemilihan Data (Data Selection)', level=2)
add_mixed_paragraph([
    ("Tahap pertama dalam proses ", False, False),
    ("Data Mining", False, True),
    (" adalah pemilihan dan pemuatan data mentah ke dalam sistem. Dataset heart_disease_health_indicators_BRFSS2015.csv berhasil dimuat ke dalam ", False, False),
    ("DataFrame", False, True),
    (" Pandas menggunakan fungsi pd.read_csv(). Hasil pemuatan data menunjukkan bahwa dataset terdiri atas ", False, False),
    ("253.680 baris", True, False),
    (" dan ", False, False),
    ("22 kolom", True, False),
    (", yang sepenuhnya sesuai dengan spesifikasi yang telah dideskripsikan pada Bab III.", False, False),
])
add_paragraph("Proses preview data menggunakan fungsi df.head() menampilkan lima baris pertama dari dataset, yang mengonfirmasi bahwa seluruh kolom telah terbaca dengan tipe data numerik (float64). Hal ini menandakan bahwa dataset telah dalam kondisi yang siap untuk diproses lebih lanjut tanpa memerlukan konversi tipe data tambahan.")

doc.add_heading('4.2 Analisis Data Eksploratif (EDA)', level=2)
doc.add_heading('4.2.1 Verifikasi Data Kosong (Missing Values)', level=3)
add_paragraph("Pemeriksaan data kosong dilakukan menggunakan fungsi df.isnull().sum() pada seluruh 22 kolom. Hasil pemeriksaan menunjukkan bahwa tidak ditemukan satu pun sel data yang kosong (null) di seluruh kolom. Dengan kata lain, dataset CDC BRFSS 2015 ini terbukti sangat bersih dan utuh, sehingga tidak diperlukan teknik pengisian data palsu (data imputation) yang berpotensi mengurangi keakuratan model.")

doc.add_heading('4.2.2 Analisis Distribusi Kelas Target', level=3)
add_mixed_paragraph([
    ("Visualisasi distribusi kelas target HeartDiseaseorAttack menggunakan grafik batang (", False, False),
    ("bar chart", False, True),
    (") melalui pustaka Seaborn mengungkapkan kondisi yang sangat signifikan: jumlah data pasien dengan label ", False, False),
    ("0 (Tidak Berisiko)", True, False),
    (" jauh lebih mendominasi dibandingkan data pasien dengan label ", False, False),
    ("1 (Berisiko)", True, False),
    (". Kondisi ini secara teknis disebut sebagai ", False, False),
    ("Imbalanced Data", False, True),
    (" (data tidak seimbang), yang merupakan fenomena umum dan alami dalam dataset medis dunia nyata, mengingat mayoritas populasi umum memang tergolong sehat.", False, False),
])

doc.add_heading('4.2.3 Analisis Korelasi Antarfitur (Heatmap)', level=3)
add_paragraph("Visualisasi heatmap korelasi menggunakan matriks korelasi Pearson menampilkan hubungan linear antara setiap pasangan fitur dalam dataset. Temuan utama dari analisis heatmap ini adalah:")
add_bullet_item([("Korelasi positif tertinggi ", True, False), ("terhadap kolom target HeartDiseaseorAttack dimiliki oleh fitur GenHlth (r ≈ 0,26), Age (r ≈ 0,22), HighBP (r ≈ 0,21), dan DiffWalk (r ≈ 0,20).", False, False)])
add_bullet_item([("Korelasi negatif ", True, False), ("teridentifikasi pada fitur Income (r ≈ -0,11) dan Education (r ≈ -0,08), yang mengindikasikan bahwa semakin tinggi pendapatan dan pendidikan seseorang, semakin rendah risiko penyakit jantungnya secara statistik.", False, False)])

doc.add_heading('4.3 Pra-pemrosesan Data (Preprocessing)', level=2)
add_paragraph("Tahap pra-pemrosesan data mencakup tiga proses utama yang dilakukan secara berurutan:")

doc.add_heading('4.3.1 Pemisahan Variabel', level=3)
add_paragraph("Data dipisahkan menjadi dua komponen menggunakan operasi DataFrame:")
add_bullet_item([("Variabel Independen (X): ", True, False), ("Seluruh 21 kolom fitur prediktor, diperoleh dengan menghapus kolom target menggunakan df.drop('HeartDiseaseorAttack', axis=1).", False, False)])
add_bullet_item([("Variabel Dependen (y): ", True, False), ("Kolom target HeartDiseaseorAttack saja, diperoleh melalui df['HeartDiseaseorAttack'].", False, False)])

doc.add_heading('4.3.2 Pembagian Data (Train-Test Split)', level=3)
add_paragraph("Data kemudian dibagi menjadi dua subset menggunakan fungsi train_test_split() dari pustaka Scikit-Learn dengan konfigurasi sebagai berikut:")
add_bullet_item([("Rasio Pembagian: ", True, False), ("80% untuk training set dan 20% untuk testing set", False, False)])
add_bullet_item([("Parameter random_state=42: ", True, False), ("Ditetapkan untuk memastikan reprodusibilitas hasil pembagian data", False, False)])
add_bullet_item([("Parameter stratify=y: ", True, False), ("Diterapkan untuk menjaga proporsi distribusi kelas target yang konsisten di kedua subset", False, False)])
add_paragraph("Hasil pembagian menghasilkan ±202.944 baris data latih dan ±50.736 baris data uji.")

doc.add_heading('4.3.3 Standardisasi (Standard Scaling)', level=3)
add_paragraph("Proses standardisasi diterapkan menggunakan objek StandardScaler() dari Scikit-Learn. Metode fit_transform() digunakan pada data latih untuk menghitung parameter statistik (rata-rata dan simpangan baku) sekaligus mentransformasi data tersebut. Sementara itu, metode transform() digunakan pada data uji untuk menerapkan parameter statistik yang sama dari data latih, sehingga tidak terjadi kebocoran informasi (data leakage) dari data uji ke proses pelatihan.")

doc.add_heading('4.4 Pemodelan Klasifikasi (Random Forest)', level=2)
add_paragraph("Model klasifikasi dibangun menggunakan kelas RandomForestClassifier dari pustaka Scikit-Learn dengan konfigurasi hiperparameter sebagai berikut:")
add_table(
    ["Parameter", "Nilai", "Justifikasi"],
    [
        ["n_estimators", "100", "Jumlah pohon keputusan yang memadai untuk prediksi stabil dan akurat"],
        ["random_state", "42", "Menjamin reprodusibilitas hasil pelatihan model"],
        ["class_weight", "balanced", "Memberikan bobot lebih tinggi pada kelas minoritas"],
    ]
)
add_paragraph("Proses pelatihan dilakukan dengan memanggil metode rf_model.fit(X_train_scaled, y_train), di mana model mempelajari pola-pola hubungan antara 21 fitur prediktor yang telah distandarisasi dengan label target penyakit jantung dari ±202.944 baris data latih.")

doc.add_heading('4.5 Evaluasi dan Validasi Model', level=2)
doc.add_heading('4.5.1 Akurasi Keseluruhan', level=3)
add_mixed_paragraph([
    ("Setelah proses pelatihan selesai, model diuji menggunakan data uji (", False, False),
    ("testing set", False, True),
    (") yang belum pernah dilihat oleh model sebelumnya. Hasil evaluasi menunjukkan bahwa model ", False, False),
    ("Random Forest", False, True),
    (" berhasil mencapai ", False, False),
    ("akurasi keseluruhan sebesar ±90%", True, False),
    (". Angka ini melampaui target minimum yang telah ditetapkan sebelumnya (85%) dan sejalan dengan temuan penelitian-penelitian terdahulu yang menggunakan dataset dan algoritma serupa.", False, False),
])

doc.add_heading('4.5.2 Laporan Klasifikasi (Classification Report)', level=3)
add_paragraph("Classification Report yang dicetak menyajikan metrik evaluasi detail per kelas:")
add_bullet_item([("Kelas 0 (Tidak Berisiko): ", True, False), ("Presisi, Recall, dan F1-Score secara konsisten berada di kisaran tinggi, menunjukkan bahwa model sangat mahir dalam mengidentifikasi pola-pola pasien yang sehat.", False, False)])
add_bullet_item([("Kelas 1 (Berisiko): ", True, False), ("Meskipun metrik evaluasinya relatif lebih rendah dibandingkan kelas 0, hal ini merupakan konsekuensi yang wajar dan terdokumentasi (expected trade-off) dari sifat alami data yang tidak seimbang (imbalanced).", False, False)])

doc.add_heading('4.5.3 Confusion Matrix', level=3)
add_paragraph("Visualisasi Confusion Matrix melalui heatmap berwarna biru mengonfirmasi distribusi prediksi model secara granular:")
add_bullet_item([("True Negative (Kiri Atas): ", True, False), ("Jumlah pasien sehat yang benar diprediksi sehat oleh model — merupakan angka terbesar.", False, False)])
add_bullet_item([("True Positive (Kanan Bawah): ", True, False), ("Jumlah pasien berisiko yang benar diprediksi berisiko oleh model.", False, False)])
add_bullet_item([("False Negative (Kiri Bawah): ", True, False), ("Jumlah pasien yang sebenarnya berisiko tetapi diprediksi sehat — ini merupakan kesalahan paling fatal dalam konteks medis.", False, False)])
add_bullet_item([("False Positive (Kanan Atas): ", True, False), ("Jumlah pasien yang sebenarnya sehat tetapi diprediksi berisiko — merupakan false alarm yang relatif lebih aman.", False, False)])

doc.add_heading('4.6 Pemodelan Klasterisasi (K-Means)', level=2)
add_mixed_paragraph([
    ("Sebagai pelengkap analisis prediktif, algoritma ", False, False),
    ("K-Means Clustering", False, True),
    (" diterapkan pada keseluruhan 253.680 baris data dengan parameter n_clusters=3. Berbeda dengan ", False, False),
    ("Random Forest", False, True),
    (" yang bersifat ", False, False),
    ("supervised", False, True),
    (" (memerlukan label target), ", False, False),
    ("K-Means", False, True),
    (" bersifat ", False, False),
    ("unsupervised", False, True),
    (" (tanpa label target), sehingga kolom HeartDiseaseorAttack tidak digunakan dalam proses ini.", False, False),
])
add_mixed_paragraph([
    ("Setelah proses klasterisasi selesai, setiap baris data diberikan label klaster (0, 1, atau 2) yang merepresentasikan segmen profil pasien yang ditemukan oleh algoritma secara alami. Visualisasi ", False, False),
    ("Scatter Plot", False, True),
    (" menggunakan dimensi ", False, False),
    ("Umur", True, False),
    (" (sumbu X) dan ", False, False),
    ("BMI", True, False),
    (" (sumbu Y) membuktikan bahwa algoritma berhasil memisahkan pasien ke dalam tiga kelompok warna yang berbeda. Segmentasi ini berpotensi sangat berguna bagi pihak medis untuk merancang program intervensi kesehatan yang tepat sasaran bagi masing-masing segmen pasien.", False, False),
])

doc.add_heading('4.7 Implementasi Aplikasi Web (Deployment)', level=2)
doc.add_heading('4.7.1 Ekspor Model (Model Export)', level=3)
add_paragraph("Tahap deployment dimulai dengan mengekspor dua objek kunci ke dalam format file pickle (.pkl) menggunakan pustaka Joblib:")
add_numbered_item(1, [("model_jantung_rf.pkl ", True, False), ("(±473 MB): Berisi seluruh struktur dan bobot dari 100 pohon keputusan Random Forest yang telah terlatih.", False, False)])
add_numbered_item(2, [("scaler_jantung.pkl ", True, False), ("(±1,5 KB): Berisi parameter statistik (rata-rata dan simpangan baku) dari StandardScaler yang digunakan selama proses pelatihan. File ini wajib disertakan agar data masukan dari pengguna di web mengalami transformasi yang identik dengan data latih.", False, False)])

doc.add_heading('4.7.2 Pengembangan Aplikasi Web Streamlit', level=3)
add_paragraph("Aplikasi web dikembangkan dalam satu file Python tunggal (app.py) dengan total 263 baris kode. Fitur-fitur utama yang diimplementasikan meliputi:")
add_numbered_item(1, [("Antarmuka Bertahap (Wizard-Style UI): ", True, False), ("Formulir masukan dibagi menjadi empat langkah yang dikendalikan melalui mekanisme st.session_state untuk menyimpan data pengguna antarlangkah tanpa kehilangan informasi.", False, False)])
add_numbered_item(2, [("Stepper Visual: ", True, False), ("Indikator progres visual berbentuk langkah-langkah bernomor yang menunjukkan posisi pengisian data pengguna saat ini.", False, False)])
add_numbered_item(3, [("Validasi dan Pemetaan Data: ", True, False), ("Setiap masukan pengguna yang bersifat kategorik dipetakan ke nilai numerik yang sesuai dengan skema encoding pada dataset asli melalui dictionary mapping.", False, False)])
add_numbered_item(4, [("Prediksi Real-Time: ", True, False), ("Setelah pengguna menekan tombol prediksi, data masukan ditransformasi menggunakan scaler yang telah dimuat, kemudian diteruskan ke model Random Forest untuk menghasilkan klasifikasi dan probabilitas secara instan.", False, False)])
add_numbered_item(5, [("Rekomendasi Medis Otomatis: ", True, False), ("Sistem secara otomatis menyajikan saran medis yang disesuaikan berdasarkan hasil prediksi.", False, False)])
add_numbered_item(6, [("Sidebar Informasi: ", True, False), ("Panel samping yang menyediakan informasi tentang aplikasi, algoritma yang digunakan, dan akurasi model.", False, False)])
add_numbered_item(7, [("Desain Responsif dan Estetis: ", True, False), ("Penggunaan font Poppins dari Google Fonts, ikon Font Awesome, gradien warna pada tombol, serta tata letak dua kolom yang rapi untuk menciptakan pengalaman pengguna yang premium.", False, False)])

doc.add_heading('4.7.3 Alur Kerja Teknis Prediksi', level=3)
add_paragraph("Alur kerja teknis yang terjadi di balik layar saat pengguna menekan tombol prediksi adalah sebagai berikut:")
add_numbered_item(1, [("Sistem mengumpulkan seluruh 21 nilai masukan dari st.session_state.", False, False)])
add_numbered_item(2, [("Nilai-nilai kategorik dikonversi menjadi nilai numerik melalui dictionary mapping.", False, False)])
add_numbered_item(3, [("Seluruh 21 nilai disusun ke dalam array NumPy berdimensi (1, 21).", False, False)])
add_numbered_item(4, [("Array tersebut ditransformasi menggunakan objek scaler yang telah dimuat dari file scaler_jantung.pkl.", False, False)])
add_numbered_item(5, [("Data yang telah distandarisasi diteruskan ke model Random Forest melalui metode model.predict() untuk memperoleh klasifikasi (0 atau 1) dan model.predict_proba() untuk memperoleh probabilitas.", False, False)])
add_numbered_item(6, [("Hasil prediksi disajikan kepada pengguna dalam format visual yang intuitif: ikon sukses hijau untuk hasil \"Sehat\" atau ikon peringatan merah untuk hasil \"Berisiko,\" disertai persentase keyakinan dan rekomendasi medis.", False, False)])

doc.add_heading('4.8 Pengujian Fungsional Aplikasi', level=2)
add_paragraph("Pengujian fungsional dilakukan untuk memverifikasi bahwa seluruh fitur aplikasi web berjalan sesuai dengan spesifikasi yang dirancang. Pengujian mencakup skenario-skenario berikut:")
add_table(
    ["No.", "Skenario Pengujian", "Hasil yang Diharapkan", "Status"],
    [
        ["1", "Pengguna mengisi Langkah 1 dan menekan \"Lanjut\"", "Halaman berpindah ke Langkah 2 tanpa kehilangan data", "Berhasil"],
        ["2", "Pengguna menekan \"Kembali\" di Langkah 3", "Halaman kembali ke Langkah 2 dengan data tetap tersimpan", "Berhasil"],
        ["3", "Pengguna mengisi seluruh data dan menekan \"Prediksi\"", "Sistem menampilkan klasifikasi, probabilitas, dan saran medis", "Berhasil"],
        ["4", "Masukan data pasien sehat (BMI normal, tidak merokok)", "Sistem memprediksi \"TIDAK BERISIKO\"", "Berhasil"],
        ["5", "Masukan data pasien berisiko tinggi (hipertensi, obesitas)", "Sistem memprediksi \"BERISIKO\" dengan peringatan merah", "Berhasil"],
        ["6", "File .pkl tidak ditemukan di direktori", "Sistem menampilkan pesan peringatan (mode simulasi)", "Berhasil"],
    ]
)
add_paragraph("Seluruh skenario pengujian berhasil dilalui tanpa ditemukan kesalahan (bug) yang signifikan, menunjukkan bahwa aplikasi web telah siap untuk didemonstrasikan.")

# =====================================================
# BAB V - PENUTUP
# =====================================================
doc.add_page_break()
doc.add_heading('BAB V', level=1)
doc.add_heading('PENUTUP', level=1)

doc.add_heading('5.1 Kesimpulan', level=2)
add_paragraph("Berdasarkan hasil penelitian dan pembahasan yang telah diuraikan pada bab-bab sebelumnya, maka dapat ditarik beberapa kesimpulan sebagai berikut:")

add_numbered_item(1, [
    ("Pra-pemrosesan data ", True, False),
    ("terhadap dataset CDC BRFSS 2015 yang berskala masif (253.680 baris dan 22 kolom) berhasil dilakukan secara optimal. Dataset terbukti bersih tanpa ", False, False),
    ("missing values", False, True),
    (", sehingga tidak diperlukan teknik imputasi. Proses standardisasi menggunakan ", False, False),
    ("StandardScaler", False, True),
    (" dan penanganan ", False, False),
    ("imbalanced data", False, True),
    (" melalui parameter class_weight='balanced' memastikan bahwa data siap untuk diproses oleh algoritma ", False, False),
    ("Machine Learning", False, True),
    (" secara akurat dan adil terhadap kedua kelas target.", False, False),
])

add_numbered_item(2, [
    ("Algoritma ", False, False),
    ("Random Forest Classifier ", True, True),
    ("berhasil mencapai ", False, False),
    ("akurasi keseluruhan sebesar ±90% ", True, False),
    ("dalam memprediksi risiko penyakit jantung. Angka ini melampaui target minimum yang telah ditetapkan (85%) dan konsisten dengan temuan penelitian-penelitian terdahulu. Analisis korelasi melalui ", False, False),
    ("heatmap", False, True),
    (" mengonfirmasi bahwa faktor-faktor seperti kondisi kesehatan umum, umur, tekanan darah tinggi, dan bahkan faktor sosioekonomi (pendidikan dan pendapatan) memiliki hubungan yang terukur dan signifikan terhadap risiko penyakit jantung.", False, False),
])

add_numbered_item(3, [
    ("Algoritma ", False, False),
    ("K-Means Clustering ", True, True),
    ("berhasil mengelompokkan 253.680 data pasien ke dalam ", False, False),
    ("3 klaster (segmen) ", True, False),
    ("yang secara visual terbukti terpisah dengan baik pada dimensi Umur dan BMI. Hasil ini membuktikan bahwa pola-pola pengelompokan alami memang ada di dalam data rekam medis dan dapat dipetakan secara akurat tanpa campur tangan label manual manusia.", False, False),
])

add_numbered_item(4, [
    ("Implementasi aplikasi web ", True, False),
    ("berbasis ", False, False),
    ("framework", False, True),
    (" Streamlit berhasil mewujudkan integrasi penuh antara model ", False, False),
    ("Machine Learning", False, True),
    (" (file .pkl) dengan antarmuka pengguna yang interaktif, estetis, dan ramah pengguna. Aplikasi ini mampu menerima masukan 21 variabel melalui formulir bertahap (", False, False),
    ("wizard-style", False, True),
    ("), mengeksekusi prediksi ", False, False),
    ("Random Forest", False, True),
    (" secara ", False, False),
    ("real-time", False, True),
    (", dan menghasilkan output berupa klasifikasi risiko, persentase keyakinan algoritma, serta rekomendasi medis otomatis. Seluruh skenario pengujian fungsional telah berhasil dilalui tanpa kesalahan.", False, False),
])

doc.add_heading('5.2 Saran', level=2)
add_paragraph("Untuk pengembangan dan penyempurnaan lebih lanjut di masa mendatang, penulis menyarankan hal-hal sebagai berikut:")
add_numbered_item(1, [("Perbandingan Multialgorima: ", True, False), ("Membandingkan kinerja Random Forest dengan algoritma-algoritma klasifikasi lain seperti Gradient Boosting (XGBoost, LightGBM), Support Vector Machine (SVM), atau Deep Learning (Jaringan Saraf Tiruan) untuk menentukan algoritma yang paling optimal.", False, False)])
add_numbered_item(2, [("Optimasi Hiperparameter: ", True, False), ("Menerapkan teknik pencarian hiperparameter otomatis seperti Grid Search atau Random Search dengan validasi silang (Cross-Validation) untuk menemukan kombinasi parameter yang lebih optimal.", False, False)])
add_numbered_item(3, [("Deployment ke Cloud: ", True, False), ("Melakukan deployment aplikasi web ke platform cloud seperti Streamlit Cloud, Heroku, atau AWS agar aplikasi dapat diakses secara publik melalui internet.", False, False)])
add_numbered_item(4, [("Pembaruan Dataset: ", True, False), ("Menggunakan data BRFSS terbaru (misalnya tahun 2020 atau 2023) yang mungkin memiliki distribusi yang lebih relevan dengan kondisi kesehatan masyarakat masa kini.", False, False)])
add_numbered_item(5, [("Penambahan Fitur Visualisasi: ", True, False), ("Menambahkan fitur visualisasi feature importance (tingkat kepentingan setiap fitur) dan SHAP values pada halaman web agar pengguna dapat memahami faktor-faktor yang paling memengaruhi hasil prediksinya secara personal.", False, False)])
add_numbered_item(6, [("Integrasi Basis Data: ", True, False), ("Menambahkan fitur penyimpanan riwayat prediksi pasien ke dalam basis data (misalnya SQLite atau PostgreSQL) agar data historis dapat dilacak dan dianalisis secara longitudinal.", False, False)])

# =====================================================
# DAFTAR PUSTAKA
# =====================================================
doc.add_page_break()
doc.add_heading('DAFTAR PUSTAKA', level=1)
add_empty_lines(1)

daftar_pustaka = [
    "Alaa, A. M., Bolton, T., Di Angelantonio, E., Rudd, J. H. F., & van der Schaar, M. (2019). Cardiovascular Disease Risk Prediction Using Automated Machine Learning: A Prospective Study of Half a Million Adults. PLOS ONE, 14(5), e0213653.",
    "Braveman, P. A., Egerter, S. A., & Williams, D. R. (2011). The Social Determinants of Health: Coming of Age. Annual Review of Public Health, 32, 381–398.",
    "Breiman, L. (2001). Random Forests. Machine Learning, 45(1), 5–32.",
    "Centers for Disease Control and Prevention [CDC]. (2015). Behavioral Risk Factor Surveillance System (BRFSS) 2015 Survey Data and Documentation. Diakses dari https://www.cdc.gov/brfss/",
    "Chawla, N. V. (2010). Data Mining for Imbalanced Datasets: An Overview. Dalam O. Maimon & L. Rokach (Eds.), Data Mining and Knowledge Discovery Handbook (pp. 875–886). Springer.",
    "Dinh, A., Miertschin, S., Young, A., & Mohanty, S. D. (2019). A Data-Driven Approach to Predicting Diabetes and Cardiovascular Disease with Machine Learning. BMC Medical Informatics and Decision Making, 19(1), 211.",
    "Fayyad, U., Piatetsky-Shapiro, G., & Smyth, P. (1996). From Data Mining to Knowledge Discovery in Databases. AI Magazine, 17(3), 37–54.",
    "Géron, A. (2019). Hands-On Machine Learning with Scikit-Learn, Keras, and TensorFlow (2nd ed.). O'Reilly Media.",
    "Han, J., Kamber, M., & Pei, J. (2012). Data Mining: Concepts and Techniques (3rd ed.). Morgan Kaufmann.",
    "Hastie, T., Tibshirani, R., & Friedman, J. (2009). The Elements of Statistical Learning: Data Mining, Inference, and Prediction (2nd ed.). Springer.",
    "Larose, D. T., & Larose, C. D. (2014). Discovering Knowledge in Data: An Introduction to Data Mining (2nd ed.). John Wiley & Sons.",
    "MacQueen, J. (1967). Some Methods for Classification and Analysis of Multivariate Observations. Dalam Proceedings of the Fifth Berkeley Symposium on Mathematical Statistics and Probability (pp. 281–297).",
    "Mohan, S., Thirumalai, C., & Srivastava, G. (2019). Effective Heart Disease Prediction Using Hybrid Machine Learning Techniques. IEEE Access, 7, 81542–81554.",
    "Pal, M., & Parija, S. (2021). Prediction of Heart Diseases Using Random Forest. Journal of Physics: Conference Series, 1817(1), 012009.",
    "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., ... & Duchesnay, É. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12, 2825–2830.",
    "Reddy, K. V. V., Elamvazuthi, I., Aziz, A. A., Paramasivam, S., Chua, H. N., & Pranavanand, S. (2020). Heart Disease Risk Prediction Using Machine Learning Classifiers with Attribute Evaluators. Applied Sciences, 10(22), 8137.",
    "Sokolova, M., & Lapalme, G. (2009). A Systematic Analysis of Performance Measures for Classification Tasks. Information Processing & Management, 45(4), 427–437.",
    "Streamlit Inc. (2024). Streamlit Documentation. Diakses dari https://docs.streamlit.io/",
    "VanderPlas, J. (2016). Python Data Science Handbook: Essential Tools for Working with Data. O'Reilly Media.",
    "World Health Organization [WHO]. (2021). Cardiovascular Diseases (CVDs). Diakses dari https://www.who.int/news-room/fact-sheets/detail/cardiovascular-diseases-(cvds)",
    "Zheng, A., & Casari, A. (2018). Feature Engineering for Machine Learning: Principles and Techniques for Data Scientists. O'Reilly Media.",
]

for ref in daftar_pustaka:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-1.25)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# ===== SIMPAN FILE =====
output_path = r'd:\KULIAH\SP SEMESTER 6\PROJECT AKHIR\Laporan_Proyek_Akhir.docx'
doc.save(output_path)
print(f"BERHASIL! File Word tersimpan di: {output_path}")
